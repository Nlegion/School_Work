# В письме обычно есть три части: приветствие, собственно письмо и прощание. Вводится текст, на основе которого нужно создать документ в формате docx, в котором первая из введенных строк – приветствие, две последние – прощание, остальные – рядовые абзацы.
# Форматирование приветствия: выравнивание по правому краю, курсив, шрифт Arial, кегль 11.
# Форматирование рядовых абзацев: выравнивание по ширине, первая строка с отступом 0.5 дюйма, шрифт Times New Roman, кегль 12.
# Форматирование прощания: выравнивание по левому краю без отступа, последняя строка – полужирное начертание, шрифт по умолчанию.
# Сохраните письмо в файл letter.docx.
# Если ввести такой текст:
# Dear Kate and Nick,
# We are looking forward very much to your visit to our country this summer. We are expecting you at the beginning of July and are hoping that you may stay until the end of the month or longer.
# We consider it a privilege for us to receive you as guests in our house. We are very grateful indeed to you for consenting to come and stay with us. We are looking forward to oﬀering you hospitality in return for the hospitality you have kindly given us on many occasions.
# We want you to understand that we will see to all your needs while you are with us and to any expenses that may arise.
# Yours sincerely,
# John and Mary

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, Mm
import sys

data = list(map(lambda x: x.replace('\n', ''), sys.stdin.readlines()))

document = Document()

p1 = document.add_paragraph(data[0])
p1.style = document.styles.add_style('P', WD_STYLE_TYPE.PARAGRAPH)
font = p1.style.font
font.name = 'Arial'
font.size = Pt(11)
font.italic = True
p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT

style = document.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

for i in data[1:-2]:
    p = document.add_paragraph(i)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_fmt = p.paragraph_format
    p_fmt.first_line_indent = Inches(0.5)

p2 = document.add_paragraph(data[-2])
p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
fmt = p2.paragraph_format
fmt.space_after = Mm(0)
p3 = document.add_paragraph(data[-1])
p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
p3.style = document.styles.add_style('P3', WD_STYLE_TYPE.PARAGRAPH)
font = p3.style.font
font.bold = True

document.save('letter.docx')
